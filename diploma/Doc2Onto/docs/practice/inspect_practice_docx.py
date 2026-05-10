"""
Утилита: показывает структуру docx-документа практики, выделяя места,
помеченные жёлтой подсветкой (это пользовательские поля для заполнения).
"""
from __future__ import annotations

import sys
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


def is_yellow(run) -> bool:
    color = run.font.highlight_color
    return color is not None and color == WD_COLOR_INDEX.YELLOW


def inspect(path: Path) -> List[str]:
    out: List[str] = []
    out.append("")
    out.append("=" * 80)
    out.append(f"FILE: {path.name}")
    out.append("=" * 80)

    doc = Document(str(path))

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        yellow_runs = [r.text for r in para.runs if is_yellow(r) and r.text.strip()]
        marker = "  <YELLOW>" if yellow_runs else ""
        if text or yellow_runs:
            out.append(f"[P{i}] {text}{marker}")
            for yr in yellow_runs:
                out.append(f"      -> жёлтый: «{yr}»")

    for ti, table in enumerate(doc.tables):
        out.append("")
        out.append(f"--- TABLE {ti} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell_text = "\n".join(p.text for p in cell.paragraphs).strip()
                yellow_in_cell = []
                for p in cell.paragraphs:
                    for r in p.runs:
                        if is_yellow(r) and r.text.strip():
                            yellow_in_cell.append(r.text)
                marker = "  <YELLOW>" if yellow_in_cell else ""
                if cell_text or yellow_in_cell:
                    out.append(f"  [R{ri}C{ci}] {cell_text}{marker}")
                    for yr in yellow_in_cell:
                        out.append(f"          -> жёлтый: «{yr}»")
    return out


if __name__ == "__main__":
    args = sys.argv[1:]
    output_path = Path(args[0]) if args and args[0].endswith(".txt") else None
    paths_args = args[1:] if output_path else args
    paths = [Path(p) for p in paths_args] or sorted(
        (Path(__file__).resolve().parent.parent / "docs" / "practice").glob("*.docx")
    )
    all_lines: List[str] = []
    for p in paths:
        all_lines.extend(inspect(p))
    if output_path:
        output_path.write_text("\n".join(all_lines), encoding="utf-8")
        print(f"Written {output_path} ({len(all_lines)} lines)")
    else:
        for line in all_lines:
            print(line)
