from pathlib import Path
from docx import Document
from docx.document import Document as DocumentObject
from lxml import etree  # type: ignore


LDML_XSD_PATH = "./modules/converter/ldml.xsd"
LDML_NS = "https://example.org/ldml/1.0"
NSMAP = {"ld": LDML_NS}


def qname(tag: str) -> str:
    """
    Формирует qualified name для элемента в namespace LDML.

    В XML элемент будет сериализован как ld:tag,
    но внутренне QName задаётся через {namespace}tag.
    """
    return f"{{{LDML_NS}}}{tag}"


class DocxToLdmlConverter:
    """
    Конвертер из DOCX в LDML.

    Текущие возможности:
    - извлечение логической структуры (текст, таблицы)
    - объединение w:r на уровне абзацев
    - валидация результата по XSD
    """

    def __init__(self):
        self.ldml_xsd_path = Path(LDML_XSD_PATH).resolve()
        self._schema = self._load_schema()

    def convert(self, input_path: str | Path) -> etree._ElementTree:
        input_path = Path(input_path)
        document = Document(str(input_path))

        # Создание корневого элемента (тег "body")
        body = etree.Element(qname("body"), nsmap=NSMAP)

        # Построение LDML-дерева по документу
        self._process_block_items(document, body)

        tree = etree.ElementTree(body)
        self._validate(tree)
        return tree

    def _load_schema(self) -> etree.XMLSchema:
        with open(self.ldml_xsd_path, "rb") as f:
            schema_doc = etree.XML(f.read())
        return etree.XMLSchema(schema_doc)

    def _validate(self, tree: etree._ElementTree) -> None:
        if not self._schema.validate(tree):
            error = self._schema.error_log.last_error
            raise ValueError(f"LDML validation error: {error}")

    def _process_block_items(self, document: DocumentObject, body_elem: etree._Element) -> None:
        """
        Проход по блочным элементам документа (абзацы и таблицы) 
        в порядке их следования.
        """

        for block in document.element.body.iterchildren():
            tag = etree.QName(block.tag).localname
            if tag == "p":
                self._process_paragraph(block, body_elem, document)
            elif tag == "tbl":
                self._process_table(block, body_elem, document)
            # прочие элементы игнорируются

    def _process_paragraph(self, p_elem, body_elem, document: DocumentObject) -> None:
        paragraph = self._find_paragraph(p_elem, document)

        text = self._extract_paragraph_text_preserve_case(paragraph)
        if not text or not text.strip():
            return

        text_elem = etree.SubElement(body_elem, qname("text"))
        text_elem.text = text.strip()

        # Семантические подсказки (необязательные)
        style = paragraph.style.name if paragraph.style else None
        if style and style.lower().startswith("heading"):
            text_elem.set("role", "heading")
            level = self._extract_heading_level(style)
            if level is not None:
                text_elem.set("level", str(level))

    def _find_paragraph(self, p_elem, document: DocumentObject):
        """
        Поиск соответствующего объекта Paragraph из python-docx
        по XML-элементу.
        """
        for p in document.paragraphs:
            if p._p is p_elem:
                return p
        raise RuntimeError("The paragraph was not found in the document")

    def _extract_paragraph_text_preserve_case(self, paragraph) -> str:
        """
        Извлекает текст абзаца с сохранением регистра.
        Учитывает w:caps ТОЛЬКО на уровне run.
        """

        result_parts: list[str] = []

        for run in paragraph.runs:
            text = run.text
            if not text:
                continue

            # Проверка w:caps на уровне run
            run_caps = (
                run._r.rPr is not None
                and run._r.rPr.caps is not None
            )

            if run_caps:
                text = text.upper()

            result_parts.append(text)

        return "".join(result_parts)

    def _extract_heading_level(self, style_name: str):
        parts = style_name.split()
        if len(parts) == 2 and parts[1].isdigit():
            return int(parts[1])
        return None

    def _process_table(self, tbl_elem, body_elem, document: DocumentObject) -> None:
        table = self._find_table(tbl_elem, document)

        table_elem = etree.SubElement(body_elem, qname("table"))

        for row in table.rows:
            row_elem = etree.SubElement(table_elem, qname("row"))
            for cell in row.cells:
                cell_elem = etree.SubElement(row_elem, qname("cell"))
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if text and text.strip():
                        text_elem = etree.SubElement(cell_elem, qname("text"))
                        text_elem.text = text

    def _find_table(self, tbl_elem, document: DocumentObject):
        """
        Поиск соответствующей таблицы python-docx по XML-элементу.
        """
        for t in document.tables:
            if t._tbl is tbl_elem:
                return t
        raise RuntimeError("The table was not found in the document")
