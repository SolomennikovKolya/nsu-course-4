from pathlib import Path
from typing import Dict, Optional, Tuple

import mammoth
from docx import Document as DocxDocument
from html import escape
from PySide6.QtCore import QObject, QRunnable, Qt, QThreadPool, QUrl, Signal
from PySide6.QtGui import QDesktopServices
from PySide6.QtWidgets import (
    QFrame,
    QLabel,
    QPushButton,
    QStackedWidget,
    QTextBrowser,
    QVBoxLayout,
    QWidget,
)

from core.document import Document
from ui.common.design import UI_COLOR_RED

try:
    from PySide6.QtPdf import QPdfDocument
    from PySide6.QtPdfWidgets import QPdfView

    _QT_PDF = True
except ImportError:
    QPdfDocument = None  # type: ignore[misc, assignment]
    QPdfView = None  # type: ignore[misc, assignment]
    _QT_PDF = False


def _docx_to_html_fragment(path: Path) -> str:
    with path.open("rb") as f:
        result = mammoth.convert_to_html(f)
    return result.value


def _wrap_original_html(body_fragment: str) -> str:
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8"/>
<style>
  body {{
    font-family: "Segoe UI", "Microsoft YaHei", "Helvetica Neue", sans-serif;
    color: #ffffff;
  }}
  body, p, li, td, th, div, span, strong, em {{ color: #ffffff; }}
  a {{ color: #a8d4ff; }}
  table {{ border-collapse: collapse; margin: 0.5em 0; }}
  td, th {{ border: 1px solid #888888; padding: 4px 6px; vertical-align: top; }}
  p {{ margin: 0.25em 0; }}
</style>
</head>
<body>{body_fragment}</body>
</html>"""


class DocumentViewOriginalTab(QWidget):
    """Вкладка для отображения оригинального документа."""

    def __init__(self):
        super().__init__()
        self._document: Optional[Document] = None
        self._pool = QThreadPool.globalInstance()
        self._request_id = 0
        self._cache: Dict[Tuple[str, int], Tuple[str, str]] = {}

        self._toolbar = QWidget()
        toolbar_layout = QVBoxLayout(self._toolbar)
        toolbar_layout.setContentsMargins(0, 0, 0, 8)

        disclaimer = QLabel("Предпросмотр может отличаться от реального вида документа")
        disclaimer.setWordWrap(True)
        disclaimer.setStyleSheet("QLabel { color: #8a8a8a; }")

        self._open_external_btn = QPushButton()
        self._open_external_btn.clicked.connect(self._open_original_externally)

        toolbar_layout.addWidget(disclaimer)
        toolbar_layout.addWidget(self._open_external_btn, alignment=Qt.AlignmentFlag.AlignLeft)

        self._stack = QStackedWidget()
        self._view = QTextBrowser()
        self._view.setReadOnly(True)
        self._view.setOpenExternalLinks(True)
        self._view.setFrameShape(QFrame.Shape.NoFrame)
        self._stack.addWidget(self._view)

        self._pdf_view: Optional["QPdfView"] = None
        self._pdf_doc: Optional["QPdfDocument"] = None

        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        layout.addWidget(self._toolbar)
        layout.addWidget(self._stack, 1)
        self._toolbar.setVisible(False)

    def set_document(self, document: Optional[Document]) -> bool:
        self._document = document
        self._view.clear()
        self._view.setPlaceholderText("")
        self._toolbar.setVisible(False)
        self._stack.setCurrentWidget(self._view)
        self._release_pdf()

        path = self._original_path()
        if path is None:
            self._view.setPlaceholderText("Документ не выбран")
            return False
        if not path.exists():
            self._view.setPlaceholderText("Оригинальный файл документа отсутствует")
            return False

        ext = path.suffix.lower()

        if ext == ".pdf":
            self._toolbar.setVisible(True)
            self._open_external_btn.setText("Открыть в программе по умолчанию")
            self._show_pdf(path)
            return True

        if ext == ".docx":
            self._toolbar.setVisible(True)
            self._open_external_btn.setText("Открыть в Word")
            self._start_docx_preview(path)
            return True

        if ext == ".doc":
            self._toolbar.setVisible(True)
            self._open_external_btn.setText("Открыть в Word")
            sidecar = path.with_suffix(".docx")
            if sidecar.is_file():
                self._start_docx_preview(sidecar)
            else:
                self._view.setPlainText(
                    "Предпросмотр .doc появится после конвертации в .docx "
                    "(рядом с файлом будет создана копия с расширением .docx). "
                    "Запустите обработку документа или откройте файл во внешнем приложении."
                )
            return True

        self._view.setPlainText(
            f"Предпросмотр для формата «{ext or '(нет расширения)'}» не поддерживается."
        )
        return True

    def _original_path(self) -> Optional[Path]:
        if self._document is None:
            return None
        return self._document.original_file_path()

    def _start_docx_preview(self, path: Path):
        cache_key = self._cache_key(path)
        cached = self._cache.get(cache_key) if cache_key is not None else None
        if cached is not None:
            mode, content = cached
            self._apply_preview(mode, content)
            return

        self._view.setPlainText("Загрузка предпросмотра…")
        self._request_id += 1
        req_id = self._request_id
        worker = _DocxPreviewWorker(req_id=req_id, path=path)
        worker.signals.finished.connect(self._on_preview_ready)
        self._pool.start(worker)

    def _show_pdf(self, path: Path):
        if not _QT_PDF:
            self._view.setPlainText(
                "Встроенный просмотр PDF недоступен: в сборке PySide6 нет модулей QtPdf / QtPdfWidgets."
            )
            return

        # setDocument(nullptr) в QPdfView в ряде версий Qt вызывает предупреждение
        # QObject::connect(QPdfDocument, QPdfLinkModel): invalid nullptr parameter.
        # Поэтому не отвязываем документ вручную: каждый показ — новый QPdfView, старый
        # убираем из стека и deleteLater (QPdfDocument — ребёнок вида, уничтожается с ним).
        if self._pdf_view is not None:
            self._dispose_pdf_preview()

        self._pdf_view = QPdfView()
        self._stack.addWidget(self._pdf_view)
        self._pdf_doc = QPdfDocument(self._pdf_view)
        self._pdf_view.setDocument(self._pdf_doc)
        self._pdf_doc.load(str(path.resolve()))
        if self._pdf_doc.pageCount() < 1:
            self._dispose_pdf_preview()
            self._view.setPlainText("Не удалось открыть PDF для предпросмотра.")
            self._stack.setCurrentWidget(self._view)
            return

        self._stack.setCurrentWidget(self._pdf_view)

    def _release_pdf(self):
        self._dispose_pdf_preview()

    def _dispose_pdf_preview(self) -> None:
        if self._pdf_view is None:
            return
        self._stack.removeWidget(self._pdf_view)
        self._pdf_view.deleteLater()
        self._pdf_view = None
        self._pdf_doc = None

        if not _QT_PDF:
            return

        # Предупреждение может прийти при отложенном удалении QPdfView; обрабатываем
        # DeferredDelete сразу и подавляем только известный шум QtPdf.
        from PySide6.QtCore import QEvent, QCoreApplication, QtMsgType, qInstallMessageHandler

        chain_to = [None]

        def _chain(msg_type, context, message: str) -> None:
            if (
                msg_type == QtMsgType.QtWarningMsg
                and "QPdfLinkModel" in message
                and "invalid nullptr parameter" in message
            ):
                return
            prev = chain_to[0]
            if prev is not None:
                prev(msg_type, context, message)

        chain_to[0] = qInstallMessageHandler(_chain)
        try:
            QCoreApplication.sendPostedEvents(None, QEvent.Type.DeferredDelete)
        finally:
            qInstallMessageHandler(chain_to[0])

    def _open_original_externally(self):
        path = self._original_path()
        if path is None or not path.exists():
            return
        QDesktopServices.openUrl(QUrl.fromLocalFile(str(path.resolve())))

    def _cache_key(self, path: Path) -> Optional[Tuple[str, int]]:
        try:
            st = path.stat()
        except OSError:
            return None
        return str(path.resolve()), int(getattr(st, "st_mtime_ns", int(st.st_mtime * 1e9)))

    def _on_preview_ready(self, req_id: int, mode: str, content: str, cache_key: Optional[Tuple[str, int]]):
        if req_id != self._request_id:
            return
        if cache_key is not None:
            self._cache[cache_key] = (mode, content)
        self._apply_preview(mode, content)

    def _apply_preview(self, mode: str, content: str):
        self._stack.setCurrentWidget(self._view)
        if mode == "html":
            self._view.setHtml(content)
        else:
            self._view.setPlainText(content)


class _DocxPreviewSignals(QObject):
    finished = Signal(int, str, str, object)


class _DocxPreviewWorker(QRunnable):
    def __init__(self, req_id: int, path: Path):
        super().__init__()
        self.req_id = req_id
        self.path = path
        self.signals = _DocxPreviewSignals()

    def run(self):
        cache_key: Optional[Tuple[str, int]]
        try:
            st = self.path.stat()
            cache_key = (str(self.path.resolve()), int(getattr(st, "st_mtime_ns", int(st.st_mtime * 1e9))))
        except OSError:
            cache_key = None

        try:
            fragment = _docx_to_html_fragment(self.path)
            html = _wrap_original_html(fragment)
            self.signals.finished.emit(self.req_id, "html", html, cache_key)
            return
        except Exception as exc:
            try:
                docx = DocxDocument(str(self.path))
                lines = [p.text for p in docx.paragraphs]
                plain = "\n".join(lines)
                note = (
                    "Не удалось показать форматированный предпросмотр "
                    f"({exc}). Ниже — только текст параграфов.\n\n"
                )
                self.signals.finished.emit(self.req_id, "plain", note + plain, cache_key)
                return
            except Exception as exc2:
                html = (
                    f"<p style='color:{UI_COLOR_RED};'>Не удалось открыть DOCX.</p>"
                    f"<pre style='color:#eeeeee;'>{escape(str(exc2))}</pre>"
                )
                self.signals.finished.emit(self.req_id, "html", html, cache_key)
